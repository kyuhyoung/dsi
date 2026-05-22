#!/usr/bin/env python3
"""마크다운 + style.yaml → 시각 디자인된 .docx 빌더.

용법:
    python3 scripts/yaml_to_docx.py <입력.md> <출력.docx> [--style <style.yaml>]

원리:
    md_to_docx.py 의 마크다운 파서를 상위 호환하면서
    style.yaml 의 *셀 fill·헤더·footer·페이지번호·표지·그림 placeholder* 모두 적용.

알고리즘만 PY. *데이터·색·임계·키워드*는 style.yaml.
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Optional

import yaml
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ──────────────────────────────────────────────
# style.yaml + form.yaml 로딩 + 기본값
# ──────────────────────────────────────────────
STYLE: dict = {}
FORM_TABLES: list = []        # form.yaml 의 tables list (cell.visual 포함)
FORM_TABLE_CURSOR = {"i": 0}  # 현재 렌더링 중인 form 표 인덱스


def load_style(path: Optional[Path]) -> dict:
    """style.yaml 로딩. 없으면 빈 dict."""
    if path is None or not path.exists():
        return {}
    with path.open(encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def load_form(path: Optional[Path]) -> list:
    """form.yaml 로딩 후 tables list 반환. 없으면 [].

    *시각 정보 (cell.visual.background_color·bold) 를 빌더가 직접 적용*하도록 사용.
    """
    if path is None or not path.exists():
        return []
    with path.open(encoding="utf-8") as f:
        data = yaml.safe_load(f) or {}
    return data.get("tables", []) or []


def form_cell_visual(table_idx: int, row: int, col: int) -> dict:
    """form.yaml 의 (table_idx, row, col) 셀의 visual dict 반환. 없으면 {}.

    form table 의 cells 는 [{row, col, text, colspan, rowspan, visual?}] list.
    """
    if not FORM_TABLES or table_idx >= len(FORM_TABLES):
        return {}
    cells = FORM_TABLES[table_idx].get("cells", [])
    for c in cells:
        if c.get("row") == row and c.get("col") == col:
            return c.get("visual", {}) or {}
    return {}


def s_get(*keys, default=None):
    """STYLE dict 안전 접근. s_get('fonts', 'body_pt', default=10)."""
    cur = STYLE
    for k in keys:
        if not isinstance(cur, dict) or k not in cur:
            return default
        cur = cur[k]
    return cur


def hex_to_rgb(h: str) -> RGBColor:
    h = (h or "").lstrip("#")
    if len(h) != 6:
        return RGBColor(0, 0, 0)
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def luminance(hex_color: str) -> float:
    """배경 색의 상대 명도 계산 (0~1). WCAG 근사. 짙으면 0에 가까움."""
    h = (hex_color or "").lstrip("#")
    if len(h) != 6:
        return 1.0
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return (0.299 * r + 0.587 * g + 0.114 * b) / 255.0


def auto_text_color(bg_hex: str) -> str:
    """배경 명도 기반 대비 텍스트 색 자동 결정. 짙은 배경 → 흰색, 밝은 배경 → 검정."""
    return "FFFFFF" if luminance(bg_hex) < 0.5 else "000000"


# ──────────────────────────────────────────────
# 메타 추출
# ──────────────────────────────────────────────
def extract_meta_from_md(md_text: str) -> dict:
    """마크다운 상단 frontmatter 또는 첫 H1에서 메타 추출."""
    meta: dict = {}
    fm = re.match(r"^---\n(.*?)\n---\n", md_text, re.DOTALL)
    if fm:
        try:
            meta.update(yaml.safe_load(fm.group(1)) or {})
        except yaml.YAMLError:
            pass
    h1 = re.search(r"^#\s+(.+)$", md_text, re.MULTILINE)
    if h1 and "사업명" not in meta:
        meta["사업명"] = h1.group(1).strip()
    return meta


def resolve_value(spec: str, meta: dict) -> str:
    """style yaml의 value_from 표현 해석. 'meta.사업명' → meta['사업명']."""
    if not isinstance(spec, str):
        return ""
    if spec.startswith("meta."):
        key = spec[len("meta."):]
        return str(meta.get(key, ""))
    return spec


# ──────────────────────────────────────────────
# 셀 fill (python-docx에 직접 API 없음, oxml로)
# ──────────────────────────────────────────────
def set_cell_fill(cell, hex_color: str):
    if not hex_color:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def set_cell_borders(cell, pt: float, color: str = "000000", bottom_pt: Optional[float] = None):
    """셀 4면 border 두께·색 지정. bottom_pt 주어지면 하단만 다른 두께."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    sz = str(int(pt * 8))   # docx border size 단위
    sz_bottom = str(int((bottom_pt or pt) * 8))
    for side in ("top", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        borders.append(b)
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), sz_bottom)
    b.set(qn("w:color"), color)
    borders.append(b)
    tc_pr.append(borders)


# ──────────────────────────────────────────────
# 페이지 설정·헤더·풋터
# ──────────────────────────────────────────────
def apply_page_setup(doc: Document):
    page = s_get("page", default={}) or {}
    margin = page.get("margin_cm") or {}
    for section in doc.sections:
        if "top" in margin:
            section.top_margin = Cm(margin["top"])
        if "bottom" in margin:
            section.bottom_margin = Cm(margin["bottom"])
        if "left" in margin:
            section.left_margin = Cm(margin["left"])
        if "right" in margin:
            section.right_margin = Cm(margin["right"])


def apply_header_footer(doc: Document, meta: dict):
    """헤더(사업명) + 풋터(페이지 번호) 적용."""
    header_cfg = s_get("header_text", default={}) or {}
    footer_cfg = s_get("footer", default={}) or {}

    for section in doc.sections:
        # 헤더
        if header_cfg.get("enabled"):
            header = section.header
            p = header.paragraphs[0]
            p.text = resolve_value(header_cfg.get("text_from", ""), meta)
            align = header_cfg.get("alignment", "right")
            p.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }.get(align, WD_ALIGN_PARAGRAPH.RIGHT)
            for run in p.runs:
                run.font.name = s_get("fonts", "default_family", default="맑은 고딕")
                run.font.size = Pt(header_cfg.get("font_pt", 9))
                run.font.color.rgb = hex_to_rgb(header_cfg.get("color", "666666"))

        # 풋터 (페이지 번호)
        if footer_cfg.get("enabled"):
            footer = section.footer
            p = footer.paragraphs[0]
            p.text = ""
            align = footer_cfg.get("alignment", "right")
            p.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }.get(align, WD_ALIGN_PARAGRAPH.RIGHT)

            fmt = footer_cfg.get("format", "{page} / {total}")
            # 페이지 번호 fields 삽입 (PAGE, NUMPAGES)
            insert_page_field(p, fmt, footer_cfg)


def insert_page_field(paragraph, fmt: str, cfg: dict):
    """{page}·{total} 자리에 docx PAGE·NUMPAGES field 삽입."""
    font_name = s_get("fonts", "default_family", default="맑은 고딕")
    font_pt = cfg.get("font_pt", 9)
    color = cfg.get("color", "666666")

    parts = re.split(r"(\{page\}|\{total\})", fmt)
    for part in parts:
        if part == "{page}":
            add_field(paragraph, "PAGE", font_name, font_pt, color)
        elif part == "{total}":
            add_field(paragraph, "NUMPAGES", font_name, font_pt, color)
        else:
            run = paragraph.add_run(part)
            run.font.name = font_name
            run.font.size = Pt(font_pt)
            run.font.color.rgb = hex_to_rgb(color)


def add_field(paragraph, field_code: str, font_name: str, font_pt: float, color: str):
    """docx field 코드 삽입 (PAGE·NUMPAGES·SEQ 등)."""
    run = paragraph.add_run()
    run.font.name = font_name
    run.font.size = Pt(font_pt)
    run.font.color.rgb = hex_to_rgb(color)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = f" {field_code} "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


# ──────────────────────────────────────────────
# 표지 페이지
# ──────────────────────────────────────────────
def add_cover_page(doc: Document, meta: dict):
    cfg = s_get("cover", default={}) or {}
    if not cfg.get("enabled"):
        return

    font_name = s_get("fonts", "default_family", default="맑은 고딕")

    # 빈 단락으로 위 여백 — 너무 많으면 빈 페이지 생성됨, 3개 정도
    for _ in range(3):
        doc.add_paragraph()

    # 제목
    title_cfg = cfg.get("title", {}) or {}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(resolve_value(title_cfg.get("text_from", "meta.사업명"), meta))
    run.font.name = font_name
    run.font.size = Pt(title_cfg.get("font_pt", 24))
    run.bold = title_cfg.get("bold", True)
    run.font.color.rgb = hex_to_rgb(title_cfg.get("color", "1F4E79"))

    # 부제
    doc.add_paragraph()
    sub_cfg = cfg.get("subtitle", {}) or {}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(sub_cfg.get("text", "사 업 계 획 서"))
    run.font.name = font_name
    run.font.size = Pt(sub_cfg.get("font_pt", 18))
    run.font.color.rgb = hex_to_rgb(sub_cfg.get("color", "404040"))

    # 메타 표 — 위 여백
    for _ in range(2):
        doc.add_paragraph()
    meta_cfg = cfg.get("meta_table", {}) or {}
    rows = meta_cfg.get("rows", []) or []
    if rows:
        table = doc.add_table(rows=len(rows), cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, row_def in enumerate(rows):
            label_cell = table.cell(i, 0)
            value_cell = table.cell(i, 1)
            label_cell.text = row_def.get("label", "")
            value_cell.text = resolve_value(row_def.get("value_from", ""), meta)
            set_cell_fill(label_cell, s_get("tables", "cell_fill", "header", default="D9D9D9"))
            for cell in (label_cell, value_cell):
                set_cell_borders(cell, s_get("tables", "borders", "all_pt", default=0.5))
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = font_name
                        r.font.size = Pt(s_get("fonts", "body_pt", default=10))

    # 표지 다음 본문 시작 — md의 <!-- pagebreak --> 가 break 제공
    # 여기서 add_page_break() 하면 빈 페이지 생김


# ──────────────────────────────────────────────
# 표 렌더링 (셀 fill + 헤더 인식 + 강조 패턴)
# ──────────────────────────────────────────────
def render_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    padded = [list(r) + [""] * (n_cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(padded), cols=n_cols)
    table.style = s_get("tables", "default_style", default="Light Grid Accent 1")

    font_name = s_get("fonts", "default_family", default="맑은 고딕")
    header_pt = s_get("fonts", "table_header_pt", default=10)
    body_pt = s_get("fonts", "table_body_pt", default=9)
    header_keywords = s_get("tables", "header_keywords", default=[]) or []
    emphasis_patterns = s_get("tables", "emphasis_patterns", default=[]) or []
    border_pt = s_get("tables", "borders", "all_pt", default=0.5)
    border_color = s_get("tables", "borders", "color", default="000000")
    header_bottom_pt = s_get("tables", "borders", "header_bottom_pt", default=1.0)

    fill_header = s_get("tables", "cell_fill", "header", default="D9D9D9")
    fill_emphasis = s_get("tables", "cell_fill", "emphasis", default="FFF2CC")

    # form.yaml 매칭 — 현재 md 표 idx 와 form.yaml 의 같은 idx 표를 매칭
    # cursor 는 (a) parse_md 의 "표 N" 헤딩 hint 또는 (b) 자동 증가
    form_table_idx = FORM_TABLE_CURSOR["i"] if FORM_TABLES else -1
    use_form_visual = 0 <= form_table_idx < len(FORM_TABLES)
    FORM_TABLE_CURSOR["i"] += 1  # 다음 표가 hint 없으면 자동 증가

    # 헤더 행 인식 — 첫 행에 header_keywords 중 하나가 있으면 헤더
    first_row_texts = [c.strip() for c in padded[0]]
    is_header_first_row = any(
        any(kw in t for kw in header_keywords) for t in first_row_texts
    )

    for i, row in enumerate(padded):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = (cell_text or "").strip()
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # *form.yaml 시각 정보 우선 적용* (있으면)
            visual = form_cell_visual(form_table_idx, i, j) if use_form_visual else {}
            bg_from_form = visual.get("background_color")
            bold_from_form = visual.get("bold", False)

            # fallback: style.yaml 기반 분류 (form.yaml 없거나 visual 없는 셀)
            is_header = (i == 0 and is_header_first_row)
            is_emphasis = any(
                re.search(pat, cell_text or "") for pat in emphasis_patterns
            )

            # fill 결정 — form.yaml visual 우선
            if bg_from_form:
                set_cell_fill(cell, bg_from_form.lstrip("#"))
                set_cell_borders(cell, border_pt, border_color)
            elif is_header:
                set_cell_fill(cell, fill_header)
                set_cell_borders(cell, border_pt, border_color, bottom_pt=header_bottom_pt)
            elif is_emphasis:
                set_cell_fill(cell, fill_emphasis)
                set_cell_borders(cell, border_pt, border_color)
            else:
                set_cell_borders(cell, border_pt, border_color)

            # 텍스트 서체 — bold: form > is_header
            apply_bold = bold_from_form or is_header
            # text 색 — form 명시 > 배경 명도 기반 자동 대비
            explicit_tc = visual.get("text_color") if visual else None
            auto_tc = auto_text_color(bg_from_form) if bg_from_form else None
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = font_name
                    if apply_bold:
                        run.font.size = Pt(header_pt)
                        run.bold = True
                    else:
                        run.font.size = Pt(body_pt)
                    if explicit_tc:
                        run.font.color.rgb = hex_to_rgb(explicit_tc)
                    elif auto_tc:
                        run.font.color.rgb = hex_to_rgb(auto_tc)


# ──────────────────────────────────────────────
# 그림 placeholder
# ──────────────────────────────────────────────
FIGURE_COUNTER = {"n": 0}
TABLE_COUNTER = {"n": 0}


def add_figure_placeholder(doc: Document, caption: str, file_hint: str):
    """그림 자리표시자 박스 + 캡션."""
    fig_cfg = s_get("figures", default={}) or {}
    style = fig_cfg.get("placeholder_style", {}) or {}

    # 자동 삽입 시도 (KB assets)
    auto = fig_cfg.get("auto_insert", {}) or {}
    inserted = False
    if auto.get("enabled"):
        asset_dir = Path(auto.get("asset_dir", ""))
        mapping = auto.get("mapping", {}) or {}
        mapped = mapping.get(file_hint)
        if mapped:
            asset_path = asset_dir / mapped
            if asset_path.exists():
                doc.add_picture(str(asset_path), width=Cm(style.get("width_cm", 14)))
                inserted = True

    if not inserted:
        # placeholder 박스 (1x1 table) — 회색 fill + 아이콘 + 라벨
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = t.cell(0, 0)
        cell.width = Cm(style.get("width_cm", 14))
        set_cell_fill(cell, style.get("fill", "F2F2F2"))
        set_cell_borders(
            cell,
            style.get("border_pt", 0.5),
            style.get("border_color", "BFBFBF"),
        )
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        icon = style.get("icon_emoji", "🖼")
        label_pt = style.get("label_font_pt", 10)
        run = p.add_run(f"{icon}  [{file_hint}]")
        run.font.name = s_get("fonts", "default_family", default="맑은 고딕")
        run.font.size = Pt(label_pt)
        run.italic = style.get("label_italic", True)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 캡션
    cap_cfg = fig_cfg.get("caption", {}) or {}
    if cap_cfg.get("auto_numbering", True):
        FIGURE_COUNTER["n"] += 1
        prefix = cap_cfg.get("prefix", "그림")
        caption_text = f"<{prefix} {FIGURE_COUNTER['n']}> {caption}"
    else:
        caption_text = caption
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption_text)
    run.font.name = s_get("fonts", "default_family", default="맑은 고딕")
    run.font.size = Pt(cap_cfg.get("font_pt", 9))
    run.italic = cap_cfg.get("italic", True)
    run.font.color.rgb = hex_to_rgb(cap_cfg.get("color", "404040"))


# ──────────────────────────────────────────────
# 마크다운 파싱 (md_to_docx.py 의 파서 확장)
# ──────────────────────────────────────────────
def parse_md(md_text: str) -> list:
    # frontmatter 제거
    md_text = re.sub(r"^---\n.*?\n---\n", "", md_text, count=1, flags=re.DOTALL)

    fig_pattern = s_get(
        "figures",
        "placeholder_pattern",
        default=r"!\[(.*?)\]\(placeholder:\s*(.+?)\)",
    )

    lines = md_text.split("\n")
    tokens = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 그림 placeholder
        m = re.match(fig_pattern, stripped)
        if m:
            tokens.append(("figure", m.group(1).strip(), m.group(2).strip()))
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            # 헤딩에 "표 N." / "표 N " / "표N" 패턴이 있으면 form table idx hint 추가
            heading_text = m.group(2)
            hint_m = re.match(r"^표\s*(\d+)[\.\s]", heading_text)
            if hint_m:
                tokens.append(("table_idx_hint", int(hint_m.group(1))))
            tokens.append(("heading", len(m.group(1)), heading_text))
            i += 1
            continue

        if stripped.startswith("```"):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            tokens.append(("code", "\n".join(buf)))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s\-:|]+\|$", lines[i + 1].strip()
        ):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if re.match(r"^\|[\s\-:|]+\|$", row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.split("|")[1:-1]]
                table_rows.append(cells)
                i += 1
            tokens.append(("table", table_rows))
            continue

        if re.match(r"^[-*]\s+", stripped):
            tokens.append(("li", stripped[2:].strip()))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            tokens.append(("oli", re.sub(r"^\d+\.\s+", "", stripped)))
            i += 1
            continue

        if stripped.startswith(">"):
            tokens.append(("quote", stripped.lstrip("> ").strip()))
            i += 1
            continue

        if stripped == "---":
            tokens.append(("hr",))
            i += 1
            continue

        if re.match(r"^<!--\s*page[\s_-]?break\s*-->$", stripped, re.IGNORECASE):
            tokens.append(("pagebreak",))
            i += 1
            continue

        if stripped:
            tokens.append(("p", stripped))
        else:
            tokens.append(("blank",))
        i += 1
    return tokens


# ──────────────────────────────────────────────
# 렌더링
# ──────────────────────────────────────────────
def add_heading(doc: Document, text: str, level: int):
    font_name = s_get("fonts", "default_family", default="맑은 고딕")
    pt_map = {
        1: s_get("fonts", "heading1_pt", default=16),
        2: s_get("fonts", "heading2_pt", default=13),
        3: s_get("fonts", "heading3_pt", default=11),
    }
    h = doc.add_heading(text.strip(), level=min(level, 4))
    for run in h.runs:
        run.font.name = font_name
        run.font.size = Pt(pt_map.get(level, s_get("fonts", "body_pt", default=10)))
    return h


def add_paragraph_md(doc: Document, text: str):
    if not text.strip():
        return
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    font_name = s_get("fonts", "default_family", default="맑은 고딕")
    body_pt = s_get("fonts", "body_pt", default=10)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.name = font_name
        run.font.size = Pt(body_pt)


def render(doc: Document, tokens: list):
    font_name = s_get("fonts", "default_family", default="맑은 고딕")
    body_pt = s_get("fonts", "body_pt", default=10)

    for token in tokens:
        t = token[0]
        if t == "heading":
            add_heading(doc, token[2], token[1])
        elif t == "p":
            add_paragraph_md(doc, token[1])
        elif t == "li":
            p = doc.add_paragraph(token[1], style="List Bullet")
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(body_pt)
        elif t == "oli":
            p = doc.add_paragraph(token[1], style="List Number")
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(body_pt)
        elif t == "table_idx_hint":
            FORM_TABLE_CURSOR["i"] = token[1]
        elif t == "table":
            render_table(doc, token[1])
        elif t == "figure":
            add_figure_placeholder(doc, token[1], token[2])
        elif t == "code":
            p = doc.add_paragraph()
            run = p.add_run(token[1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif t == "quote":
            p = doc.add_paragraph(token[1])
            for run in p.runs:
                run.italic = True
                run.font.name = font_name
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif t == "hr":
            doc.add_paragraph("─" * 40)
        elif t == "pagebreak":
            doc.add_page_break()
        elif t == "blank":
            pass


# ──────────────────────────────────────────────
# main
# ──────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("src", help="입력 .md")
    parser.add_argument("dst", help="출력 .docx")
    parser.add_argument("--style", help="style.yaml 경로", default=None)
    parser.add_argument("--form", help="form.yaml 경로 (셀 visual 매칭)", default=None)
    args = parser.parse_args()

    global STYLE, FORM_TABLES
    style_path = Path(args.style) if args.style else None
    STYLE = load_style(style_path)
    form_path = Path(args.form) if args.form else None
    FORM_TABLES = load_form(form_path)

    src = Path(args.src)
    dst = Path(args.dst)
    md_text = src.read_text(encoding="utf-8")
    meta = extract_meta_from_md(md_text)

    doc = Document()
    # 기본 폰트
    normal = doc.styles["Normal"]
    normal.font.name = s_get("fonts", "default_family", default="맑은 고딕")
    normal.font.size = Pt(s_get("fonts", "body_pt", default=10))

    apply_page_setup(doc)
    apply_header_footer(doc, meta)

    if s_get("cover", "enabled", default=False):
        add_cover_page(doc, meta)

    tokens = parse_md(md_text)
    render(doc, tokens)

    doc.save(dst)
    print(f"saved: {dst}", file=sys.stderr)
    print(
        f"  figures: {FIGURE_COUNTER['n']}, "
        f"form_tables: {len(FORM_TABLES)}, "
        f"style: {args.style or 'none'}, form: {args.form or 'none'}",
        file=sys.stderr,
    )


if __name__ == "__main__":
    main()
