#!/usr/bin/env python3
"""Markdown → docx 변환 (간이판). 헤딩·문단·표·리스트·코드블록 지원."""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level):
    h = doc.add_heading(text.strip(), level=level)
    for run in h.runs:
        run.font.name = "맑은 고딕"
    return h


def add_paragraph(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.name = "맑은 고딕"
        run.font.size = Pt(10)
    return p


def add_table(doc, rows):
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text.strip()
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "맑은 고딕"
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def parse_md(md_text):
    """라인 단위 토큰화."""
    lines = md_text.split("\n")
    tokens = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            tokens.append(("heading", level, m.group(2)))
            i += 1
            continue

        if stripped.startswith("```"):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            tokens.append(("code", "\n".join(buf)))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if re.match(r"^\|[\s\-:|]+\|$", row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.split("|")[1:-1]]
                table_rows.append(cells)
                i += 1
            tokens.append(("table", table_rows))
            continue

        if re.match(r"^[-*]\s+", stripped):
            tokens.append(("li", stripped[2:].strip()))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            tokens.append(("oli", re.sub(r"^\d+\.\s+", "", stripped)))
            i += 1
            continue

        if stripped.startswith(">"):
            tokens.append(("quote", stripped.lstrip("> ").strip()))
            i += 1
            continue

        if stripped == "---":
            tokens.append(("hr",))
            i += 1
            continue

        if stripped:
            tokens.append(("p", stripped))
        else:
            tokens.append(("blank",))
        i += 1
    return tokens


def render(doc, tokens):
    for token in tokens:
        t = token[0]
        if t == "heading":
            add_heading(doc, token[2], min(token[1], 4))
        elif t == "p":
            add_paragraph(doc, token[1])
        elif t == "li":
            p = doc.add_paragraph(token[1], style="List Bullet")
            for run in p.runs:
                run.font.name = "맑은 고딕"
                run.font.size = Pt(10)
        elif t == "oli":
            p = doc.add_paragraph(token[1], style="List Number")
            for run in p.runs:
                run.font.name = "맑은 고딕"
                run.font.size = Pt(10)
        elif t == "table":
            add_table(doc, token[1])
        elif t == "code":
            add_code_block(doc, token[1])
        elif t == "quote":
            p = doc.add_paragraph(token[1])
            for run in p.runs:
                run.italic = True
                run.font.name = "맑은 고딕"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif t == "hr":
            doc.add_paragraph("─" * 40)
        elif t == "blank":
            pass


def main():
    src = sys.argv[1]
    dst = sys.argv[2]
    md_text = Path(src).read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    tokens = parse_md(md_text)
    render(doc, tokens)
    doc.save(dst)
    print(f"saved: {dst}")


if __name__ == "__main__":
    main()
