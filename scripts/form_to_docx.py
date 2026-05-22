#!/usr/bin/env python3
"""form.yaml + filled_cells.yaml → docx 직접 빌더.

용법:
    python3 scripts/form_to_docx.py <form.yaml> <filled_cells.yaml> <out.docx> \
        [--style <style.yaml>]

원리:
    *셀 단위 흐름의 완성*. md 우회.
    - form.yaml: 양식의 표·셀 구조 + 시각 정보 (background_color·bold) — extract_hwp_form.py 산출
    - filled_cells.yaml: 셀별 (table_idx, row, col, text, source?) — proposal-writer 산출
    - style.yaml: 폰트·페이지·헤더/풋터 (옵션)

    *모든 표를 form.yaml 그대로 재현*. filled_cells 가 텍스트 채움.
    빈 셀 (filled_cells 에 매칭 없음) 은 form.yaml의 placeholder 텍스트 그대로 유지 또는 공란.

알고리즘만 PY. 데이터·색은 yaml.
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Optional

import yaml
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ──────────────────────────────────────────────
# yaml 로딩 + 기본값
# ──────────────────────────────────────────────
STYLE: dict = {}


def s_get(*keys, default=None):
    cur = STYLE
    for k in keys:
        if not isinstance(cur, dict) or k not in cur:
            return default
        cur = cur[k]
    return cur


def load_yaml(path: Path) -> dict:
    with path.open(encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


# ──────────────────────────────────────────────
# 색 유틸
# ──────────────────────────────────────────────
def hex_to_rgb(h: str) -> RGBColor:
    h = (h or "").lstrip("#")
    if len(h) != 6:
        return RGBColor(0, 0, 0)
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def luminance(hex_color: str) -> float:
    h = (hex_color or "").lstrip("#")
    if len(h) != 6:
        return 1.0
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return (0.299 * r + 0.587 * g + 0.114 * b) / 255.0


def auto_text_color(bg_hex: str) -> str:
    return "FFFFFF" if luminance(bg_hex) < 0.5 else "000000"


# ──────────────────────────────────────────────
# 셀 fill·border (oxml)
# ──────────────────────────────────────────────
def set_cell_fill(cell, hex_color: str):
    if not hex_color:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def set_cell_borders(cell, pt: float = 0.5, color: str = "000000"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    sz = str(int(pt * 8))
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


# ──────────────────────────────────────────────
# 페이지·헤더·풋터
# ──────────────────────────────────────────────
def apply_page_setup(doc: Document):
    page = s_get("page", default={}) or {}
    margin = page.get("margin_cm") or {}
    for section in doc.sections:
        if "top" in margin:
            section.top_margin = Cm(margin["top"])
        if "bottom" in margin:
            section.bottom_margin = Cm(margin["bottom"])
        if "left" in margin:
            section.left_margin = Cm(margin["left"])
        if "right" in margin:
            section.right_margin = Cm(margin["right"])


def add_field(paragraph, field_code: str, font_name: str, font_pt: float, color: str):
    run = paragraph.add_run()
    run.font.name = font_name
    run.font.size = Pt(font_pt)
    run.font.color.rgb = hex_to_rgb(color)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = f" {field_code} "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def apply_header_footer(doc: Document, meta: dict):
    header_cfg = s_get("header_text", default={}) or {}
    footer_cfg = s_get("footer", default={}) or {}
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    font_name = s_get("fonts", "default_family", default="맑은 고딕")

    for section in doc.sections:
        if header_cfg.get("enabled"):
            p = section.header.paragraphs[0]
            spec = header_cfg.get("text_from", "")
            if spec.startswith("meta."):
                p.text = str(meta.get(spec[5:], ""))
            else:
                p.text = spec
            p.alignment = align_map.get(header_cfg.get("alignment", "right"), WD_ALIGN_PARAGRAPH.RIGHT)
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(header_cfg.get("font_pt", 9))
                run.font.color.rgb = hex_to_rgb(header_cfg.get("color", "666666"))

        if footer_cfg.get("enabled"):
            p = section.footer.paragraphs[0]
            p.text = ""
            p.alignment = align_map.get(footer_cfg.get("alignment", "right"), WD_ALIGN_PARAGRAPH.RIGHT)
            fmt = footer_cfg.get("format", "{page} / {total}")
            import re
            parts = re.split(r"(\{page\}|\{total\})", fmt)
            pt = footer_cfg.get("font_pt", 9)
            color = footer_cfg.get("color", "666666")
            for part in parts:
                if part == "{page}":
                    add_field(p, "PAGE", font_name, pt, color)
                elif part == "{total}":
                    add_field(p, "NUMPAGES", font_name, pt, color)
                else:
                    r = p.add_run(part)
                    r.font.name = font_name
                    r.font.size = Pt(pt)
                    r.font.color.rgb = hex_to_rgb(color)


# ──────────────────────────────────────────────
# 핵심: 표 렌더링 (form.yaml + filled_cells.yaml 매칭)
# ──────────────────────────────────────────────
def render_form_table(doc: Document, form_table: dict, filled_by_pos: dict):
    """form.yaml 의 한 표 + filled_by_pos[(row,col)] → docx table.

    form.yaml: rows·cols·cells[{row, col, text, colspan, rowspan, visual?}]
    filled_by_pos: {(row, col): {text, source}}
    """
    rows = form_table.get("rows", 0)
    cols = form_table.get("cols", 0)
    cells = form_table.get("cells", [])
    if rows == 0 or cols == 0:
        return

    table = doc.add_table(rows=rows, cols=cols)
    table.style = s_get("tables", "default_style", default="Light Grid Accent 1")

    font_name = s_get("fonts", "default_family", default="맑은 고딕")
    body_pt = s_get("fonts", "table_body_pt", default=9)
    header_pt = s_get("fonts", "table_header_pt", default=10)
    border_pt = s_get("tables", "borders", "all_pt", default=0.5)
    border_color = s_get("tables", "borders", "color", default="000000")

    # cells 는 {row, col} key 로 조회 가능하게 dict로
    form_by_pos = {(c.get("row"), c.get("col")): c for c in cells}

    for r in range(rows):
        for col in range(cols):
            cell = table.cell(r, col)
            form_cell = form_by_pos.get((r, col), {})
            filled = filled_by_pos.get((r, col), {})

            # 텍스트: filled 우선, 없으면 form 의 원본 텍스트 (placeholder)
            text = filled.get("text") if filled.get("text") else form_cell.get("text", "")
            cell.text = str(text or "").strip()
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # 시각: form.yaml 의 visual 적용
            visual = form_cell.get("visual", {}) or {}
            bg = visual.get("background_color")
            bold = visual.get("bold", False)
            explicit_tc = visual.get("text_color")
            italic = visual.get("italic", False)

            if bg:
                set_cell_fill(cell, bg.lstrip("#"))
            set_cell_borders(cell, border_pt, border_color)

            auto_tc = auto_text_color(bg) if bg else None
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = font_name
                    run.font.size = Pt(header_pt if bold else body_pt)
                    if bold:
                        run.bold = True
                    if italic:
                        run.italic = True
                    if explicit_tc:
                        run.font.color.rgb = hex_to_rgb(explicit_tc)
                    elif auto_tc:
                        run.font.color.rgb = hex_to_rgb(auto_tc)


# ──────────────────────────────────────────────
# main 흐름
# ──────────────────────────────────────────────
def build(form_data: dict, filled_data: dict, out_path: Path):
    tables = form_data.get("tables", [])
    page_breaks_at_paragraph = set(form_data.get("page_breaks_at_paragraph", []) or [])

    # filled_cells.yaml 의 구조 가정:
    # {
    #   meta: {사업명, 신청기관, ...},
    #   table_order: [13, 14, 15, ..., 0, 1, 2, ...],  # 본체 우선 큐레이션 옵션
    #   page_break_after_table: [16, 23, ...],          # 표 N 다음 page break (옵션)
    #   filled_cells: [
    #     {table_idx: 14, row: 1, col: 1, text: '주식회사 다비오', source: 'kb/...'},
    #     ...
    #   ]
    # }
    meta = filled_data.get("meta", {}) or {}
    table_order = filled_data.get("table_order")
    if not table_order:
        # 기본: form.yaml 표 순서대로
        table_order = [t["idx"] for t in tables]
    page_break_after = set(filled_data.get("page_break_after_table", []) or [])

    # filled_cells 를 table_idx 별로 그룹화
    filled_by_table: dict = {}
    for fc in filled_data.get("filled_cells", []) or []:
        t_idx = fc.get("table_idx")
        if t_idx is None:
            continue
        pos = (fc.get("row"), fc.get("col"))
        filled_by_table.setdefault(t_idx, {})[pos] = fc

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = s_get("fonts", "default_family", default="맑은 고딕")
    normal.font.size = Pt(s_get("fonts", "body_pt", default=10))

    apply_page_setup(doc)
    apply_header_footer(doc, meta)

    # form.yaml 의 table_idx → table dict 매핑
    table_by_idx = {t["idx"]: t for t in tables}

    rendered = 0
    for t_idx in table_order:
        ft = table_by_idx.get(t_idx)
        if ft is None:
            continue
        render_form_table(doc, ft, filled_by_table.get(t_idx, {}))
        rendered += 1
        # 표 사이 짧은 공백 단락
        doc.add_paragraph()
        if t_idx in page_break_after:
            doc.add_page_break()

    doc.save(out_path)
    return rendered


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("form", help="form.yaml 경로")
    parser.add_argument("filled", help="filled_cells.yaml 경로")
    parser.add_argument("out", help="출력 .docx")
    parser.add_argument("--style", default=None)
    args = parser.parse_args()

    global STYLE
    if args.style:
        STYLE = load_yaml(Path(args.style))

    form_data = load_yaml(Path(args.form))
    filled_data = load_yaml(Path(args.filled))

    rendered = build(form_data, filled_data, Path(args.out))
    print(f"saved: {args.out}", file=sys.stderr)
    print(
        f"  form_tables: {len(form_data.get('tables', []))}, "
        f"rendered: {rendered}, "
        f"filled_cells: {len(filled_data.get('filled_cells', []) or [])}",
        file=sys.stderr,
    )


if __name__ == "__main__":
    main()
