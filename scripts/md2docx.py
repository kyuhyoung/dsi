"""제안서 마크다운 → .docx 변환기

DSI 시스템 smoke test용 간이 변환기.
실제 운영에서는 공식 docx skill의 docx-js (node) 사용 권장.

지원 마크다운:
- # ## ### #### 헤딩
- 일반 단락
- 표 | ... |
- 글머리 - 리스트
- 인용 > 블록
- ``` 코드 블록
- 수평선 ---
"""

import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_table(lines, start_idx):
    """| ... | 형식 표 파싱. 시작 인덱스부터 표 끝까지 처리.
    Returns: (rows, next_idx)
    """
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].startswith('|'):
        line = lines[i].strip()
        # 구분선 (|---|---|) 스킵
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=n_cols)
    t.style = 'Light Grid Accent 1'
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            if c_idx < n_cols:
                t.cell(r_idx, c_idx).text = cell


def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    lines = content.split('\n')
    i = 0

    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]

        # 코드 블록
        if line.startswith('```'):
            if in_code:
                # 닫는 ```
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_buf))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 수평선
        if line.strip() == '---':
            doc.add_paragraph('_' * 40)
            i += 1
            continue

        # 표
        if line.strip().startswith('|') and line.strip().endswith('|'):
            rows, next_i = parse_table(lines, i)
            add_table(doc, rows)
            i = next_i
            continue

        # 헤딩
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2), level=min(level, 9))
            i += 1
            continue

        # 인용
        if line.startswith('>'):
            p = doc.add_paragraph()
            run = p.add_run(line[1:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
            i += 1
            continue

        # 리스트
        if re.match(r'^[\s]*[-*+]\s+', line):
            text = re.sub(r'^[\s]*[-*+]\s+', '', line)
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # 번호 리스트
        if re.match(r'^[\s]*\d+\.\s+', line):
            text = re.sub(r'^[\s]*\d+\.\s+', '', line)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # 빈 줄
        if line.strip() == '':
            i += 1
            continue

        # 일반 단락
        doc.add_paragraph(line)
        i += 1

    doc.save(docx_path)
    print(f"✓ 생성: {docx_path}")


if __name__ == '__main__':
    md_to_docx(sys.argv[1], sys.argv[2])
